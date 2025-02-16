import sys
import sqlite3
from PyQt6.QtWidgets import (QWidget, QVBoxLayout, QPushButton,QLineEdit, QHBoxLayout,
                             QTableWidget, QTableWidgetItem, QTextBrowser, QDialog, QMessageBox)
from dialogs import RecordDialog, AnalysisResultDialog
from PyQt6.QtGui import *
from docx import Document

class RecordsWindow(QWidget):
    def __init__(self, medworker_id):
        super().__init__()
        self.setWindowTitle("Записи")
        self.setGeometry(100, 100, 800, 600)
        self.connection = sqlite3.connect('medical_system.db')
        self.medworker_id = medworker_id

        layout = QVBoxLayout()

        search_layout = QHBoxLayout()
        self.search_record_input = QLineEdit()
        self.search_record_input.setPlaceholderText("Поиск записи (По фамилии пациента)")
        search_layout.addWidget(self.search_record_input)

        search_button = QPushButton("Поиск")
        search_button.setIcon(QIcon('images/magnifier.png'))
        search_button.clicked.connect(self.search_record)
        search_layout.addWidget(search_button)

        search_button = QPushButton("")
        search_button.setIcon(QIcon('images/arrow-curve-180-left.png'))
        search_button.clicked.connect(self.load_records)
        search_layout.addWidget(search_button)

        layout.addLayout(search_layout)

        button_layout = QHBoxLayout()
        self.add_record_button = QPushButton("Добавить запись")
        self.add_record_button.clicked.connect(self.add_record)
        button_layout.addWidget(self.add_record_button)

        self.delete_record_button = QPushButton("Удалить запись")
        self.delete_record_button.clicked.connect(self.delete_record)
        button_layout.addWidget(self.delete_record_button)

        self.update_record_button = QPushButton("Изменить данные запись")
        self.update_record_button.clicked.connect(self.edit_record)
        button_layout.addWidget(self.update_record_button)

        layout.addLayout(button_layout)

        self.record_table = QTableWidget()
        self.record_table.setColumnCount(5)
        self.record_table.setHorizontalHeaderLabels(
            ["ID", "Пациент", "Дата записи", "Время записи", "Анализ"]
        )
        self.record_table.setSortingEnabled(True)
        layout.addWidget(self.record_table)

        self.create_record_button = QPushButton("Добавить результат")
        self.create_record_button.clicked.connect(self.create_result)
        layout.addWidget(self.create_record_button)

        self.view_records_button = QPushButton("Просмотр результата")
        self.view_records_button.clicked.connect(self.view_result)
        layout.addWidget(self.view_records_button)

        self.export_button = QPushButton("Экспорт в Word")
        self.export_button.clicked.connect(self.export_to_word)
        layout.addWidget(self.export_button)

        self.load_records()

        self.result_browser = QTextBrowser()
        layout.addWidget(self.result_browser)

        self.setLayout(layout)

    def load_records(self):
        connection = sqlite3.connect('medical_system.db')
        cursor = connection.cursor()
        cursor.execute('''
        SELECT
            pr.ID_record,
            p.Last_Name || ' ' || SUBSTR(p.First_Name, 1, 1) || '. ' || SUBSTR(p.Middle_Name, 1, 1) || '.' AS Full_Name,
            pr.Date_of_Appointment,
            pr.Time_of_Appointment,
            a.Analysis_Name
        FROM
            Patient_Records pr
        JOIN
            Patients p ON pr.ID_patient = p.ID_patient
        JOIN
            Analyses a ON pr.ID_analysis = a.ID_analysis
        ORDER BY
            pr.Date_of_Appointment, pr.Time_of_Appointment
        ''')
        rows = cursor.fetchall()
        connection.close()

        self.record_table.setRowCount(0)
        for row in rows:
            row_position = self.record_table.rowCount()
            self.record_table.insertRow(row_position)
            for column, data in enumerate(row):
                item = QTableWidgetItem(str(data))
                self.record_table.setItem(row_position, column, item)

            connection = sqlite3.connect('medical_system.db')
            cursor = connection.cursor()
            cursor.execute("SELECT * FROM Analysis_Results WHERE ID_record = ?", (row[0],))
            result = cursor.fetchone()
            connection.close()

            if result:
                for column in range(self.record_table.columnCount()):
                    self.record_table.item(row_position, column).setBackground(QColor(232, 232, 232, 127))
            else:
                for column in range(self.record_table.columnCount()):
                    self.record_table.item(row_position, column).setBackground(QColor(182, 227, 240, 127))

    def search_record(self):
        cursor = self.connection.cursor()
        search_text = self.search_record_input.text()

        if search_text:
            cursor.execute('''
        SELECT
            pr.ID_record,
            p.Last_Name || ' ' || SUBSTR(p.First_Name, 1, 1) || '. ' || SUBSTR(p.Middle_Name, 1, 1) || '.' AS Full_Name,
            pr.Date_of_Appointment,
            pr.Time_of_Appointment,
            a.Analysis_Name
        FROM
            Patient_Records pr
        JOIN
            Patients p ON pr.ID_patient = p.ID_patient
        JOIN
            Analyses a ON pr.ID_analysis = a.ID_analysis
        WHERE
            p.Last_Name LIKE ?
        ORDER BY
            pr.Date_of_Appointment, pr.Time_of_Appointment
        ''', ('%' + search_text + '%',))

            data = cursor.fetchall()
            self.record_table.setRowCount(0)

            if data:
                self.record_table.setRowCount(len(data))
                for row_index, row_data in enumerate(data):
                    for column_index, column_data in enumerate(row_data):
                        self.record_table.setItem(row_index, column_index, QTableWidgetItem(str(column_data)))
            else:
                QMessageBox.information(self, "Результат", "Записи с этой фамилией не найдены.")
        else:
            QMessageBox.warning(self, "Ошибка", "Введите текст для поиска.")

    def add_record(self):
        dialog = RecordDialog(self)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            self.load_records()

    def delete_record(self):
        row = self.record_table.currentRow()
        if row < 0:
            QMessageBox.warning(self, "Ошибка", "Выберите запись для удаления.")
            return

        record_id = self.record_table.item(row, 0).text()
        connection = sqlite3.connect('medical_system.db')
        cursor = connection.cursor()
        cursor.execute("DELETE FROM Patient_Records WHERE ID_record = ?", (record_id,))
        connection.commit()
        connection.close()
        self.load_records()

    def edit_record(self):
        row = self.record_table.currentRow()
        if row < 0:
            QMessageBox.warning(self, "Ошибка", "Выберите запись для обновления.")
            return

        record_id = self.record_table.item(row, 0).text()

        print(f"Editing record ID: {record_id}")

        try:
            dialog = RecordDialog(self, record_id=record_id)
            if dialog.exec() == QDialog.DialogCode.Accepted:
                self.load_records()
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Произошла ошибка: {str(e)}")

    def view_result(self):
        row = self.record_table.currentRow()
        if row < 0:
            QMessageBox.warning(self, "Ошибка", "Выберите запись для просмотра результата.")
            return

        record_id = self.record_table.item(row, 0).text()

        connection = sqlite3.connect('medical_system.db')
        cursor = connection.cursor()
        cursor.execute("SELECT * FROM Analysis_Results WHERE ID_record = ?", (record_id,))
        result = cursor.fetchone()

        if result:
            cursor.execute("SELECT Last_Name, First_Name, Middle_Name FROM Medworkers WHERE ID_medworker = ?",
                           (result[2],))
            executor = cursor.fetchone()
            executor_name = f"{executor[0]} {executor[1][0]}. {executor[2][0]}." if executor else "Неизвестно"

            report_text = (
                "<h2 style='text-align: center; font-weight: bold;'>Отчет о проведенном анализе</h2>"
                f"<p style='font-weight: bold;'>Дата получения результатов анализа: </p><p>{result[3]}</p>"
                f"<p style='font-weight: bold;'>Результат: </p><p>{result[4]}</p>"
                f"<p style='font-weight: bold;'>Исполнитель: </p><p>{executor_name}</p>"
            )

            self.result_browser.setHtml(report_text)
        else:
            self.result_browser.setText("Результатов ещё нет.")

        self.result_browser.setStyleSheet("font-size: 14px;")

    def create_result(self):
        row = self.record_table.currentRow()
        if row < 0:
            QMessageBox.warning(self, "Ошибка", "Выберите запись для добавления результата.")
            return

        record_id = self.record_table.item(row, 0).text()

        connection = sqlite3.connect('medical_system.db')
        cursor = connection.cursor()
        cursor.execute("SELECT * FROM Analysis_Results WHERE ID_record = ?", (record_id,))
        existing_result = cursor.fetchone()
        connection.close()

        if existing_result:
            QMessageBox.warning(self, "Ошибка", "Для этой записи уже существует результат.")
            return

        dialog = AnalysisResultDialog(self, self.medworker_id, record_id=record_id)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            self.load_records()

    def export_to_word(self):
        row = self.record_table.currentRow()
        if row < 0:
            QMessageBox.warning(self, "Ошибка", "Выберите запись для экспорта.")
            return

        record_id = self.record_table.item(row, 0).text()

        connection = sqlite3.connect('medical_system.db')
        cursor = connection.cursor()
        cursor.execute("SELECT * FROM Analysis_Results WHERE ID_record = ?", (record_id,))
        result = cursor.fetchone()

        if result:
            cursor.execute("SELECT Last_Name, First_Name, Middle_Name FROM Medworkers WHERE ID_medworker = ?",
                           (result[2],))
            executor = cursor.fetchone()
            executor_name = f"{executor[0]} {executor[1][0]}. {executor[2][0]}." if executor else "Неизвестно"

            doc = Document()
            doc.add_heading('Отчет о проведенном анализе', level=1)

            doc.add_paragraph(f"Дата получения результатов анализа: {result[3]}")
            doc.add_paragraph(f"Результат: {result[4]}")
            doc.add_paragraph(f"Исполнитель: {executor_name}")

            doc.save('report.docx')
            QMessageBox.information(self, "Успех", "Отчет успешно экспортирован в Word.")
        else:
            QMessageBox.warning(self, "Ошибка", "Результатов ещё нет.")

        connection.close()

    def analysis_result(self):
        connection = sqlite3.connect('medical_system.db')
        cursor = connection.cursor()
        cursor.execute("SELECT * FROM Analysis_Results")
        rows = cursor.fetchall()
        connection.close()

        self.record_table.setRowCount(0)
        for row in rows:
            row_position = self.record_table.rowCount()
            self.record_table.insertRow(row_position)
            for column, data in enumerate(row):
                self.record_table.setItem(row_position, column, QTableWidgetItem(str(data)))

        connection.close()
